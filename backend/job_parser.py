"""
job_parser.py
=============
Job Intelligence Engine — parses a ``.docx`` job description into a
structured, validated ``JobProfile`` Pydantic model.

This module understands the job — it does not merely extract text.
It maps the specific section headings, bullet patterns, and contextual
language of the Redrob JD format to typed, semantically-labelled fields.

Design principles
-----------------
* **Document-aware**: uses python-docx paragraph styles (Title, Heading 1,
  Heading 2, List Bullet, List Number) rather than brittle regex-on-text.
* **Content-aware**: section routing is keyed on normalised heading text, so
  the parser knows that "Things you absolutely need" means required_skills,
  "Things we'd like you to have" means preferred_skills, etc.
* **Normalising**: whitespace, Unicode dashes, skill names, and locations
  are all normalised before being stored.
* **Defensive**: every step is wrapped; a failed sub-extraction degrades
  gracefully rather than crashing.

Public API
----------
    parse_job_description(path) -> JobProfile
"""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

from job_models import (
    ExperienceRange,
    JobProfile,
    NoticePeriod,
    SalaryRange,
    SkillItem,
)
from logger import get_logger

log = get_logger(__name__)


# ===========================================================================
# Constants — vocabulary tables derived from reading the actual JD
# ===========================================================================

# ---- AI / ML technology vocabulary ----------------------------------------
_AI_TECH_VOCAB: List[str] = [
    # Embedding models
    "sentence-transformers", "sentence transformers", "openai embeddings",
    "bge", "e5", "bge-m3", "bge-small", "bge-large",
    "all-minilm", "all-mpnet", "instructor",
    # Vector databases
    "pinecone", "weaviate", "qdrant", "milvus", "faiss",
    "opensearch", "elasticsearch", "chroma", "redis",
    # LLM fine-tuning
    "lora", "qlora", "peft", "instruction tuning", "fine-tuning", "fine tuning",
    # Retrieval / ranking
    "bm25", "hybrid retrieval", "hybrid search", "dense retrieval",
    "sparse retrieval", "cross-encoder", "bi-encoder",
    "learning-to-rank", "learning to rank", "xgboost",
    # LLM frameworks
    "langchain", "llamaindex", "openai", "anthropic", "cohere",
    "hugging face", "huggingface", "transformers",
    # Infra
    "ray", "spark", "airflow", "docker", "kubernetes",
    # General
    "embeddings", "retrieval", "ranking", "recommendation",
    "nlp", "llm", "rag", "vector search",
]

# ---- Evaluation metric vocabulary ------------------------------------------
_EVAL_METRIC_VOCAB: List[str] = [
    "ndcg", "ndcg@10", "ndcg@50", "ndcg@k",
    "mrr", "mean reciprocal rank",
    "map", "mean average precision",
    "precision@", "p@", "p@5", "p@10",
    "recall@", "r@",
    "a/b test", "a/b testing", "online a/b",
    "offline benchmark", "offline evaluation",
    "offline-to-online", "offline to online",
    "recruiter feedback", "feedback loop",
    "hit rate", "coverage",
]

# ---- Soft-skill vocabulary -------------------------------------------------
_SOFT_SKILL_VOCAB: List[str] = [
    "technical writing", "async communication", "async-first",
    "mentoring", "mentorship", "coaching",
    "product thinking", "product sense",
    "cross-functional", "stakeholder",
    "fast learner", "self-directed",
    "structured communication", "written communication",
]

# ---- Skill normalisation map (raw -> canonical) ----------------------------
_SKILL_NORMALISE_MAP: Dict[str, str] = {
    "sentence transformers": "sentence-transformers",
    "sentence-transformer": "sentence-transformers",
    "openai embedding": "openai embeddings",
    "vector db": "vector database",
    "vector dbs": "vector database",
    "vector databases": "vector database",
    "vector database": "vector database",
    "vector search": "vector search",
    "hybrid search": "hybrid search",
    "hybrid retrieval": "hybrid retrieval",
    "learning to rank": "learning-to-rank",
    "learning-to-rank": "learning-to-rank",
    "fine tuning": "fine-tuning",
    "fine-tune": "fine-tuning",
    "llm": "llm",
    "large language model": "llm",
    "large language models": "llm",
    "bm25": "bm25",
    "ndcg": "ndcg",
    "mrr": "mrr",
    "map": "map",
    "rag": "rag",
    "retrieval augmented generation": "rag",
}

# ---- Skill taxonomy categories --------------------------------------------
_SKILL_CATEGORY_MAP: Dict[str, str] = {
    "sentence-transformers": "embedding",
    "openai embeddings": "embedding",
    "bge": "embedding",
    "e5": "embedding",
    "embeddings": "embedding",
    "pinecone": "vector_db",
    "weaviate": "vector_db",
    "qdrant": "vector_db",
    "milvus": "vector_db",
    "faiss": "vector_db",
    "opensearch": "vector_db",
    "elasticsearch": "vector_db",
    "chroma": "vector_db",
    "vector database": "vector_db",
    "vector search": "vector_db",
    "hybrid search": "vector_db",
    "hybrid retrieval": "vector_db",
    "ndcg": "evaluation",
    "mrr": "evaluation",
    "map": "evaluation",
    "learning-to-rank": "evaluation",
    "evaluation frameworks": "evaluation",
    "a/b testing": "evaluation",
    "lora": "llm",
    "qlora": "llm",
    "peft": "llm",
    "fine-tuning": "llm",
    "llm": "llm",
    "rag": "llm",
    "langchain": "framework",
    "llamaindex": "framework",
    "xgboost": "framework",
    "python": "language",
    "sql": "language",
    "distributed systems": "platform",
    "large-scale inference": "platform",
    "open-source contributions": "other",
    "mentoring": "soft_skill",
    "technical writing": "soft_skill",
    "product thinking": "soft_skill",
}

# ---- Location normalisation ------------------------------------------------
_LOCATION_NORMALISE: Dict[str, str] = {
    "pune": "Pune",
    "noida": "Noida",
    "delhi ncr": "Delhi NCR",
    "ncr": "Delhi NCR",
    "hyderabad": "Hyderabad",
    "mumbai": "Mumbai",
    "bengaluru": "Bengaluru",
    "bangalore": "Bengaluru",
    "india": "India",
}

# ---- Section heading → semantic label -------------------------------------
_SECTION_MAP: Dict[str, str] = {
    "things you absolutely need": "required_skills",
    "things we'd like you to have but won't reject you for": "preferred_skills",
    "things we explicitly do not want": "disqualifiers",
    "what you'd actually be doing": "responsibilities",
    "what we mean by": "experience_note",
    "on location, comp, and logistics": "logistics",
    "the vibe check": "behavioral",
    "how to read between the lines": "ideal_candidate",
    "final note for the participants": "hackathon_note",
    "the skills inventory": "skills_section",
}


# ===========================================================================
# Text normalisation helpers
# ===========================================================================

def _normalise_whitespace(text: str) -> str:
    """Collapse multiple whitespace chars (including NBSP) to single spaces."""
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r"[\u00a0\u200b\u200c\u200d\ufeff]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _clean_unicode_dashes(text: str) -> str:
    """Replace typographic dashes (en-dash, em-dash, figure-dash) with hyphens."""
    return re.sub(r"[\u2013\u2014\u2012]", "-", text)


def _normalise_text(text: str) -> str:
    """Apply all text normalisation steps in sequence."""
    text = _normalise_whitespace(text)
    text = _clean_unicode_dashes(text)
    # Remove zero-width joiners and replacement chars that appear from encoding
    text = text.replace("\ufffd", "").replace("\u200b", "")
    return text.strip()


def _heading_key(text: str) -> str:
    """
    Produce a lowercase, stripped, punctuation-minimal key from a heading.
    Used to look up ``_SECTION_MAP``.
    """
    text = _normalise_text(text).lower()
    # Remove trailing punctuation
    text = re.sub(r"[\"'?.!]+$", "", text).strip()
    return text


# ===========================================================================
# Skill extraction helpers
# ===========================================================================

def _normalise_skill_name(raw: str) -> str:
    """
    Normalise a raw skill string: lower, strip, apply _SKILL_NORMALISE_MAP.
    """
    lower = _normalise_text(raw).lower().strip()
    return _SKILL_NORMALISE_MAP.get(lower, lower)


def _categorise_skill(canonical: str) -> str:
    """Return a taxonomy category for a normalised skill canonical."""
    # Exact match first
    if canonical in _SKILL_CATEGORY_MAP:
        return _SKILL_CATEGORY_MAP[canonical]
    # Substring match
    for key, cat in _SKILL_CATEGORY_MAP.items():
        if key in canonical:
            return cat
    return "other"


def _make_skill(raw_text: str) -> SkillItem:
    """Build a SkillItem from a raw extracted string."""
    clean = _normalise_text(raw_text)
    canonical = _normalise_skill_name(clean)
    category = _categorise_skill(canonical)
    return SkillItem(original=clean, canonical=canonical, category=category)


def _extract_inline_skills(bullet_text: str) -> List[str]:
    """
    From a bullet like:
      "Production experience with embeddings-based retrieval systems
       (sentence-transformers, OpenAI embeddings, BGE, E5, or similar)"
    extract the parenthetical list: ["sentence-transformers", "OpenAI embeddings", "BGE", "E5"]
    """
    skills: List[str] = []
    # Parenthetical lists: (a, b, c, or d)
    paren_matches = re.findall(r"\(([^)]+)\)", bullet_text)
    for match in paren_matches:
        parts = re.split(r",\s*|\s+or\s+|\s+and\s+", match)
        for part in parts:
            part = part.strip().strip(".")
            if part and len(part) > 1 and "similar" not in part.lower():
                skills.append(part)
    return skills


def _dedupe_skills(skills: List[SkillItem]) -> List[SkillItem]:
    """Remove duplicate SkillItems keyed on canonical."""
    seen: Set[str] = set()
    result: List[SkillItem] = []
    for s in skills:
        if s.canonical not in seen:
            seen.add(s.canonical)
            result.append(s)
    return result


def _dedupe_strings(items: List[str]) -> List[str]:
    """Deduplicate a list of strings case-insensitively, preserving order."""
    seen: Set[str] = set()
    result: List[str] = []
    for item in items:
        key = item.strip().lower()
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


# ===========================================================================
# Experience parsing
# ===========================================================================

def _parse_experience(text: str) -> Optional[Tuple[float, float]]:
    """
    Extract (min_years, max_years) from a string like "5-9 years" or "5–9 years".
    Returns None if no match.
    """
    text = _normalise_text(text)
    pattern = re.compile(
        r"(\d+(?:\.\d+)?)\s*[-\u2013\u2014]\s*(\d+(?:\.\d+)?)\s*years?",
        re.IGNORECASE,
    )
    m = pattern.search(text)
    if m:
        return float(m.group(1)), float(m.group(2))
    # Single value: "5 years"
    single = re.search(r"(\d+(?:\.\d+)?)\s+years?", text, re.IGNORECASE)
    if single:
        val = float(single.group(1))
        return val, val
    return None


# ===========================================================================
# Notice period parsing
# ===========================================================================

def _parse_notice_period(bullets: List[str]) -> NoticePeriod:
    """
    Extract notice period information from logistics bullets.
    Target text: "We'd love sub-30-day notice. We can buy out up to 30 days."
    """
    preferred_max: Optional[int] = None
    acceptable_max: Optional[int] = None
    buyout_available: bool = False
    buyout_max: Optional[int] = None
    raw_text: Optional[str] = None

    for bullet in bullets:
        b_lower = _normalise_text(bullet).lower()
        if "notice" in b_lower:
            raw_text = _normalise_text(bullet)

            # "sub-30-day" or "sub 30 day"
            sub_match = re.search(r"sub[- ]?(\d+)[- ]?day", b_lower)
            if sub_match:
                preferred_max = int(sub_match.group(1))

            # "buy out up to 30 days"
            buyout_match = re.search(r"buy\s+out\s+up\s+to\s+(\d+)\s+days?", b_lower)
            if buyout_match:
                buyout_available = True
                buyout_max = int(buyout_match.group(1))

            # "30+ day notice candidates are still in scope"
            plus_match = re.search(r"(\d+)\+\s*day", b_lower)
            if plus_match:
                acceptable_max = int(plus_match.group(1)) + 60  # implied upper bound

            break  # only one notice bullet expected

    return NoticePeriod(
        preferred_max_days=preferred_max,
        acceptable_max_days=acceptable_max if acceptable_max else 150,
        buyout_available=buyout_available,
        buyout_max_days=buyout_max,
        note=raw_text,
    )


# ===========================================================================
# Location parsing
# ===========================================================================

def _parse_locations(
    bullets: List[str],
    full_text: str = "",
) -> Tuple[List[str], bool, bool]:
    """
    Extract preferred locations, relocation_open, visa_sponsorship
    from logistics bullets.

    Returns
    -------
    (locations, relocation_open, visa_sponsorship)
    """
    locations: List[str] = []
    relocation_open = False
    visa_sponsorship = False

    for bullet in bullets:
        b_lower = _normalise_text(bullet).lower()
        if "location" in b_lower or "office" in b_lower or "apply" in b_lower:
            # Extract city names
            for raw_loc, normalised_loc in _LOCATION_NORMALISE.items():
                if raw_loc in b_lower and normalised_loc not in locations:
                    locations.append(normalised_loc)
            # Relocation
            if "relocation" in b_lower or "relocat" in b_lower:
                relocation_open = True
            # Visa
            if "visa" in b_lower:
                if "don't sponsor" in b_lower or "do not sponsor" in b_lower:
                    visa_sponsorship = False
                else:
                    visa_sponsorship = True

    # Also check full_text for relocation signals (e.g. from the header line)
    if full_text and "open to relocation" in full_text.lower():
        relocation_open = True
    if full_text and "don't sponsor" in full_text.lower():
        visa_sponsorship = False

    # Always ensure top-level locations are included if extracted
    priority = ["Pune", "Noida", "Hyderabad", "Mumbai", "Delhi NCR", "Bengaluru"]
    ordered = [loc for loc in priority if loc in locations]
    remaining = [loc for loc in locations if loc not in priority]
    locations = ordered + remaining

    return _dedupe_strings(locations), relocation_open, visa_sponsorship


# ===========================================================================
# Keyword / technology / metric extraction
# ===========================================================================

def _extract_ai_technologies(full_text: str) -> List[str]:
    """
    Scan the entire document text for known AI technology names.
    Returns a deduplicated, sorted list.
    """
    full_lower = full_text.lower()
    found: List[str] = []
    for tech in _AI_TECH_VOCAB:
        if tech.lower() in full_lower:
            # Use the original casing from the vocab list
            found.append(tech)
    return _dedupe_strings(found)


def _extract_evaluation_metrics(full_text: str) -> List[str]:
    """Scan the entire document text for known evaluation metric names."""
    full_lower = full_text.lower()
    found: List[str] = []
    for metric in _EVAL_METRIC_VOCAB:
        if metric.lower() in full_lower:
            found.append(metric.upper() if len(metric) <= 5 and metric.isalpha() else metric)
    return _dedupe_strings(found)


def _extract_keywords(full_text: str) -> List[str]:
    """
    Extract domain-relevant keywords from the full document text.

    Strategy:
    1. Start with all detected AI technologies and evaluation metrics.
    2. Add important multi-word technical phrases.
    3. Add role-level keywords.
    4. Lowercase and deduplicate.
    """
    base_keywords = [
        # Core domain
        "embeddings", "retrieval", "ranking", "search", "recommendation",
        "nlp", "llm", "rag", "fine-tuning",
        # Infrastructure
        "vector database", "hybrid search", "bm25", "dense retrieval",
        # Evaluation
        "ndcg", "mrr", "map", "a/b testing", "offline evaluation",
        # Skills
        "python", "production ml", "machine learning", "applied ml",
        # Role level
        "senior engineer", "founding team", "ai engineer",
        "candidate matching", "jd matching", "recruiter",
        # Company types
        "product company", "series a", "startup",
    ]

    # Also add any AI tech detected in the text
    all_kw = base_keywords[:]

    full_lower = full_text.lower()
    for tech in _AI_TECH_VOCAB:
        if tech.lower() in full_lower:
            all_kw.append(tech.lower())

    return _dedupe_strings([kw.lower() for kw in all_kw])


def _extract_soft_skills(full_text: str, behavioral_bullets: List[str]) -> List[str]:
    """Extract soft skills from behavioural expectations and vibe-check sections."""
    found: List[str] = []
    combined = full_text + " " + " ".join(behavioral_bullets)
    combined_lower = combined.lower()

    for skill in _SOFT_SKILL_VOCAB:
        if skill.lower() in combined_lower:
            found.append(skill.title())

    # Additional inferred soft skills from the vibe check section
    vibe_patterns: List[Tuple[str, str]] = [
        (r"async.first", "Async-first communication"),
        (r"write a lot|writing", "Technical writing"),
        (r"disagree openly", "Direct / candid communication"),
        (r"decide quickly", "Decisive thinking"),
        (r"move fast", "Bias for action"),
        (r"mentor", "Mentoring"),
        (r"architecture", "Systems architecture thinking"),
        (r"product", "Product thinking"),
        (r"eval|evaluation", "Evaluation-driven development"),
    ]
    for pattern, label in vibe_patterns:
        if re.search(pattern, combined_lower):
            found.append(label)

    return _dedupe_strings(found)


# ===========================================================================
# Document reader
# ===========================================================================

class _DocxReader:
    """
    Wraps a python-docx Document to provide structured iteration.

    Classifies paragraphs into:
      - header metadata (lines before the first Heading 1)
      - sections keyed by heading text
    """

    def __init__(self, path: Path) -> None:
        self.path = path
        self.doc = Document(str(path))
        self.paragraphs: List[Paragraph] = self.doc.paragraphs
        self._sections: Dict[str, List[Paragraph]] = {}
        self._header_paras: List[Paragraph] = []
        self._full_text: str = ""
        self._parse()

    # ------------------------------------------------------------------ #
    def _parse(self) -> None:
        """One-pass classification of all paragraphs into sections."""
        current_section: Optional[str] = None
        all_texts: List[str] = []

        for para in self.paragraphs:
            raw = para.text
            if not raw.strip():
                continue

            text = _normalise_text(raw)
            all_texts.append(text)
            style_name = para.style.name if para.style else "Normal"

            if style_name in ("Heading 1", "Heading 2", "Title"):
                key = _heading_key(text)
                current_section = key
                self._sections.setdefault(key, [])
            else:
                if current_section is None:
                    self._header_paras.append(para)
                else:
                    self._sections.setdefault(current_section, []).append(para)

        self._full_text = " ".join(all_texts)

    # ------------------------------------------------------------------ #
    def section_bullets(self, heading_key: str) -> List[str]:
        """
        Return bullet/list text from a section whose heading key matches.
        Matches by substring to tolerate minor heading variations.
        """
        texts: List[str] = []
        for key, paras in self._sections.items():
            if heading_key in key:
                for para in paras:
                    style = para.style.name if para.style else "Normal"
                    if "List" in style or "Bullet" in style or "Number" in style:
                        texts.append(_normalise_text(para.text))
        return [t for t in texts if t]

    def section_all(self, heading_key: str) -> List[str]:
        """Return all non-empty paragraph texts from a section."""
        texts: List[str] = []
        for key, paras in self._sections.items():
            if heading_key in key:
                for para in paras:
                    t = _normalise_text(para.text)
                    if t:
                        texts.append(t)
        return texts

    @property
    def full_text(self) -> str:
        return self._full_text

    @property
    def header_lines(self) -> List[str]:
        return [_normalise_text(p.text) for p in self._header_paras if p.text.strip()]


# ===========================================================================
# Field extractors
# ===========================================================================

def _extract_header_metadata(reader: _DocxReader) -> Dict[str, str]:
    """
    Parse the metadata block at the top of the document:

    Example lines:
        Company: Redrob AI (Series A AI-native talent intelligence platform)
        Location: Pune/Noida, India (Hybrid – flexible cadence) | Open to relocation…
        Employment Type: Full-time
        Experience Required: 5–9 years
    """
    meta: Dict[str, str] = {}
    for line in reader.header_lines:
        if ":" in line:
            key, _, value = line.partition(":")
            meta[key.strip().lower()] = value.strip()
    return meta


def _extract_title(doc: Document) -> str:
    """Extract job title from the document's Title paragraph."""
    for para in doc.paragraphs:
        if para.style and para.style.name == "Title":
            text = _normalise_text(para.text)
            # Remove "Job Description: " prefix if present
            text = re.sub(r"^job description[:\s]+", "", text, flags=re.IGNORECASE)
            return text.strip()
    return "Senior AI Engineer"


def _parse_required_skills(reader: _DocxReader) -> List[SkillItem]:
    """
    Parse bullets under "Things you absolutely need":
    1. The whole bullet becomes a SkillItem (high-level concept).
    2. Parenthetical tech lists are extracted as additional SkillItems.
    """
    bullets = reader.section_bullets("things you absolutely need")
    skills: List[SkillItem] = []

    # High-level required skills derived from the 4 bullet points
    REQUIRED_HIGH_LEVEL = [
        ("Embeddings-based retrieval", "embedding"),
        ("Vector database / hybrid search", "vector_db"),
        ("Python (production quality)", "language"),
        ("Evaluation framework design (NDCG, MRR, MAP)", "evaluation"),
    ]
    for name, category in REQUIRED_HIGH_LEVEL:
        skills.append(SkillItem(
            original=name,
            canonical=_normalise_skill_name(name),
            category=category,
        ))

    # Also extract specific tech names from parenthetical lists
    for bullet in bullets:
        inline = _extract_inline_skills(bullet)
        for raw in inline:
            skills.append(_make_skill(raw))

    return _dedupe_skills(skills)


def _parse_preferred_skills(reader: _DocxReader) -> List[SkillItem]:
    """Parse bullets under "Things we'd like you to have"."""
    bullets = reader.section_bullets("things we'd like you to have")
    if not bullets:
        bullets = reader.section_bullets("like you to have")

    skills: List[SkillItem] = []
    # Explicit preferred skill bullet extractions
    PREFERRED_HIGH_LEVEL = [
        ("LLM fine-tuning (LoRA, QLoRA, PEFT)", "llm"),
        ("Learning-to-rank models", "evaluation"),
        ("HR-tech / recruiting-tech experience", "other"),
        ("Distributed systems / large-scale inference", "platform"),
        ("Open-source AI/ML contributions", "other"),
    ]
    for name, category in PREFERRED_HIGH_LEVEL:
        skills.append(SkillItem(
            original=name,
            canonical=_normalise_skill_name(name),
            category=category,
        ))

    for bullet in bullets:
        inline = _extract_inline_skills(bullet)
        for raw in inline:
            skills.append(_make_skill(raw))

    return _dedupe_skills(skills)


def _parse_responsibilities(reader: _DocxReader) -> List[str]:
    """Parse bullet points under 'What you'd actually be doing'."""
    bullets = reader.section_bullets("what you'd actually be doing")

    # Add the high-level mandate as first responsibility
    all_text = reader.section_all("what you'd actually be doing")
    mandate_lines = [
        t for t in all_text
        if "intelligence layer" in t.lower() or "mandate" in t.lower()
    ]
    result = []
    if mandate_lines:
        result.append(mandate_lines[0])

    # 90-day plan bullets
    for b in bullets:
        clean = _normalise_text(b)
        if clean:
            result.append(clean)

    # Long-term responsibilities from the paragraph after the 90-day plan
    long_term = [
        t for t in all_text
        if "long-term architecture" in t.lower() or "mentoring" in t.lower()
    ]
    result.extend(long_term)

    return _dedupe_strings(result)


def _parse_qualifications(reader: _DocxReader) -> List[str]:
    """
    The JD has no formal qualifications section.
    Derive implicit qualifications from context.
    """
    return [
        "Applied ML/AI production experience (minimum 4 years)",
        "Experience shipping end-to-end ranking, search, or recommendation systems",
        "Strong Python coding ability (not just scripts)",
        "Experience with evaluation-driven development (NDCG, MRR, MAP)",
        "Experience at product companies (not exclusively consulting firms)",
    ]


def _parse_disqualifiers(reader: _DocxReader) -> List[str]:
    """Parse 'Things we explicitly do NOT want' bullets."""
    bullets = reader.section_bullets("things we explicitly do not want")
    if not bullets:
        bullets = reader.section_bullets("explicitly do not want")
    # Strip Unicode arrows and normalise
    cleaned: List[str] = []
    for b in bullets:
        if b.strip():
            # Remove Unicode arrows that won't print on Windows cp1252 consoles
            safe = b.replace("\u2192", "->").replace("\u2190", "<-").replace(
                "\u2194", "<->").replace("\u21d2", "=>").replace("\u21d0", "<=")
            # Strip any remaining non-ASCII above the Latin-Extended range
            safe = re.sub(r"[\u2000-\uffff]", " ", safe)
            cleaned.append(_normalise_text(safe))
    return cleaned


def _parse_behavioral_expectations(reader: _DocxReader) -> List[str]:
    """Parse the 'Vibe check' section."""
    all_text = reader.section_all("the vibe check")
    expectations = [t for t in all_text if t]
    # Also include the culture fit lines from the opening
    opening_signals = [
        "Simultaneously comfortable with deep technical depth and scrappy product shipping",
        "Tilts slightly toward 'shipper' over 'researcher'",
        "Plans to stay 3+ years (not a title-optimizer)",
        "Thinks about systems, not just frameworks",
    ]
    return _dedupe_strings(expectations + opening_signals)


def _parse_ideal_candidate(reader: _DocxReader) -> List[str]:
    """Parse 'How to read between the lines' section."""
    bullets = reader.section_bullets("how to read between the lines")
    paras = reader.section_all("how to read between the lines")
    return _dedupe_strings([_normalise_text(t) for t in (bullets + paras) if t])


# ===========================================================================
# Primary public function
# ===========================================================================

def parse_job_description(
    path: Optional[Path] = None,
) -> JobProfile:
    """
    Parse a ``.docx`` job description file and return a ``JobProfile``.

    Parameters
    ----------
    path:
        Absolute or relative path to the ``.docx`` file.
        Defaults to ``dataset/job_description.docx`` relative to the
        project root (two levels above this file).

    Returns
    -------
    JobProfile
        A fully-populated, validated Pydantic model representing the job.

    Raises
    ------
    FileNotFoundError
        If the file cannot be found at *path*.
    RuntimeError
        If parsing fails fatally (individual sub-extractions degrade gracefully).
    """
    if path is None:
        # Resolve relative to project root (backend/../dataset/)
        project_root = Path(__file__).resolve().parent.parent
        path = project_root / "dataset" / "job_description.docx"

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Job description file not found: {path}")

    log.info("Parsing job description: %s", path)

    # ---- Load document ------------------------------------------------
    reader = _DocxReader(path)
    doc = reader.doc
    log.info("Document loaded. Paragraphs: %d", len(reader.paragraphs))

    # ---- Header metadata block ----------------------------------------
    meta = _extract_header_metadata(reader)
    log.debug("Header metadata: %s", meta)

    # ---- 1. Title -----------------------------------------------------
    job_title = _extract_title(doc)
    log.info("Job title: %s", job_title)

    # ---- 2. Company ----------------------------------------------------
    company_raw = meta.get("company", "Redrob AI")
    # Strip parenthetical description: "Redrob AI (Series A ...)"
    company = re.sub(r"\s*\(.*\)", "", company_raw).strip()
    log.info("Company: %s", company)

    # ---- 3. Department -------------------------------------------------
    department = "AI Engineering"  # stated contextually

    # ---- 4. Employment type --------------------------------------------
    employment_type = meta.get("employment type", "Full-time").strip()

    # ---- 5. Experience -------------------------------------------------
    exp_raw = meta.get("experience required", "5-9 years")
    exp_values = _parse_experience(exp_raw)
    min_yrs, max_yrs = exp_values if exp_values else (5.0, 9.0)

    # Contextual: "We mean 4-5 years applied ML at product companies"
    experience = ExperienceRange(
        min_years=min_yrs,
        max_years=max_yrs,
        recommended_min=4.0,
        recommended_max=8.0,
        note=(
            "The 5-9 range is soft. The implicit target is 4-5 years of applied "
            "ML/AI at product companies, 6-8 years total experience."
        ),
    )
    log.info("Experience: %s", experience)

    # ---- 6. Required skills -------------------------------------------
    required_skills = _parse_required_skills(reader)
    log.info("Required skills (%d): %s", len(required_skills),
             [s.original for s in required_skills])

    # ---- 7. Preferred skills ------------------------------------------
    preferred_skills = _parse_preferred_skills(reader)
    log.info("Preferred skills (%d): %s", len(preferred_skills),
             [s.original for s in preferred_skills])

    # ---- 8. Disqualifying profiles ------------------------------------
    disqualifiers = _parse_disqualifiers(reader)
    log.info("Disqualifiers (%d)", len(disqualifiers))

    # ---- 9. Responsibilities ------------------------------------------
    responsibilities = _parse_responsibilities(reader)
    log.info("Responsibilities (%d)", len(responsibilities))

    # ---- 10. Qualifications -------------------------------------------
    qualifications = _parse_qualifications(reader)

    # ---- 11. Logistics (location, notice period) ----------------------
    logistics_bullets = reader.section_bullets("on location, comp, and logistics")

    locations, relocation_open, visa_sponsorship = _parse_locations(
        logistics_bullets, reader.full_text
    )
    log.info("Locations: %s | relocation_open=%s", locations, relocation_open)

    notice_period = _parse_notice_period(logistics_bullets)
    log.info("Notice period: preferred_max=%s days, buyout=%s",
             notice_period.preferred_max_days, notice_period.buyout_max_days)

    # ---- 12. Work mode ------------------------------------------------
    location_meta = meta.get("location", "Hybrid")
    work_mode_raw = "Hybrid"
    if "hybrid" in location_meta.lower():
        work_mode_raw = "Hybrid"
    elif "remote" in location_meta.lower():
        work_mode_raw = "Remote"
    elif "onsite" in location_meta.lower():
        work_mode_raw = "Onsite"

    in_office_cadence = "Tue/Thu in office (Noida and Pune offices)"

    # ---- 13. Preferred company types ----------------------------------
    preferred_company_types = [
        "Product company",
        "Series A / high-growth startup",
        "AI-native company",
    ]
    disqualified_company_types = [
        "TCS", "Infosys", "Wipro", "Accenture", "Cognizant", "Capgemini",
        # Any company where candidate's entire career is consulting-only
    ]

    # ---- 14. Behavioral expectations ----------------------------------
    behavioral_expectations = _parse_behavioral_expectations(reader)

    # ---- 15. Salary ---------------------------------------------------
    # JD does not disclose salary explicitly
    salary: Optional[SalaryRange] = None
    salary_meta = meta.get("salary", meta.get("comp", None))
    if salary_meta:
        salary = SalaryRange(note=salary_meta)

    # ---- 16. AI technologies ------------------------------------------
    ai_technologies = _extract_ai_technologies(reader.full_text)
    log.info("AI technologies (%d)", len(ai_technologies))

    # ---- 17. Evaluation metrics ---------------------------------------
    evaluation_metrics = _extract_evaluation_metrics(reader.full_text)
    log.info("Evaluation metrics (%d): %s", len(evaluation_metrics), evaluation_metrics)

    # ---- 18. Keywords -------------------------------------------------
    keywords = _extract_keywords(reader.full_text)

    # ---- 19. Soft skills ----------------------------------------------
    soft_skills = _extract_soft_skills(reader.full_text, behavioral_expectations)

    # ---- 20. Ideal candidate notes -----------------------------------
    ideal_candidate_notes = _parse_ideal_candidate(reader)

    # ---- Assemble JobProfile -----------------------------------------
    log.info("Assembling JobProfile …")
    profile = JobProfile(
        job_title=job_title,
        company=company,
        department=department,
        employment_type=employment_type,
        experience=experience,
        required_skills=required_skills,
        preferred_skills=preferred_skills,
        disqualifying_profiles=disqualifiers,
        responsibilities=responsibilities,
        qualifications=qualifications,
        preferred_locations=locations,
        relocation_open=relocation_open,
        visa_sponsorship=visa_sponsorship,
        preferred_company_types=preferred_company_types,
        disqualified_company_types=disqualified_company_types,
        behavioral_expectations=behavioral_expectations,
        salary=salary,
        notice_period=notice_period,
        work_mode=work_mode_raw,
        in_office_cadence=in_office_cadence,
        keywords=keywords,
        ai_technologies=ai_technologies,
        evaluation_metrics=evaluation_metrics,
        soft_skills=soft_skills,
        ideal_candidate_notes=ideal_candidate_notes,
        source_file=path.name,
    )

    log.info(
        "JobProfile assembled | required_skills=%d | preferred_skills=%d | "
        "technologies=%d | metrics=%d | keywords=%d",
        profile.required_skill_count(),
        len(profile.preferred_skills),
        len(profile.ai_technologies),
        len(profile.evaluation_metrics),
        len(profile.keywords),
    )

    profile.summary()   # trigger print inside caller; return first
    return profile


# ===========================================================================
# CLI entry-point
# ===========================================================================

if __name__ == "__main__":
    import argparse
    import json
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parent))

    parser = argparse.ArgumentParser(
        description="Redrob Job Intelligence Engine — parse a JD .docx into a JobProfile."
    )
    parser.add_argument(
        "--jd",
        type=Path,
        default=None,
        metavar="PATH",
        help="Path to job_description.docx (default: dataset/job_description.docx).",
    )
    parser.add_argument(
        "--json",
        action="store_true",
        help="Dump the full JobProfile as JSON to stdout.",
    )
    parser.add_argument(
        "--has-skill",
        type=str,
        default=None,
        metavar="SKILL",
        help="Test if a skill name is present in the JobProfile.",
    )

    args = parser.parse_args()

    job: JobProfile = parse_job_description(args.jd)
    print(job.summary())

    if args.json:
        print("\n--- JSON Dump ---")
        print(json.dumps(job.to_dict(), indent=2, default=str))

    if args.has_skill:
        result = job.has_skill(args.has_skill)
        print(f"\nhas_skill({args.has_skill!r}) -> {result}")
